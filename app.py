# --------------------------------------------------------------
# KL Tools – Enhanced Document Management Suite (v7.5)
# --------------------------------------------------------------
# Install required packages (run once):
# pip install streamlit pandas pdfplumber openpyxl pypdf fuzzywuzzy python-Levenshtein pdf2docx python-docx Pillow img2pdf docx2pdf pywin32
# --------------------------------------------------------------

import streamlit as st
import pandas as pd
import pdfplumber
import zipfile
import io
import re
import shutil
from pathlib import Path
import tempfile
import os
import concurrent.futures
from functools import partial
from fuzzywuzzy import fuzz
from pypdf import PdfReader, PdfWriter
from pdf2docx import Converter
from docx import Document
from docx.shared import Pt
from PIL import Image                       # Pillow – image handling
import img2pdf                              # loss‑less PNG → PDF
from docx2pdf import convert                # DOCX → PDF (requires MS Word on Windows)

# ------------------------------------------------------------------
# Additional imports for the cross‑platform DOCX → PDF conversion
# ------------------------------------------------------------------
import sys
import subprocess

# --------------------------------------------------------------
# Helper for COM‑initialisation (required by docx2pdf on Windows)
# --------------------------------------------------------------
def convert_docx_to_pdf(docx_path: str, pdf_path: str) -> bool:
    """
    Convert a .docx file to .pdf.

    * Windows → uses docx2pdf (requires Microsoft Word)
    * macOS / Linux → uses LibreOffice headless conversion (requires `soffice`)

    Returns True on success, False on any error (and shows a Streamlit error widget).
    """
    try:
        # ---------------- Windows (COM) ----------------
        if sys.platform.startswith("win"):
            try:
                import pythoncom
                pythoncom.CoInitialize()
            except Exception:  # pragma: no‑cover
                pass

            # docx2pdf.convert raises an exception on failure
            convert(docx_path, pdf_path)
            return True

        # ---------------- macOS / Linux (LibreOffice) -------------
        # LibreOffice must be on the PATH as `soffice`.
        out_dir = os.path.dirname(pdf_path) or "."
        cmd = [
            "soffice",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            out_dir,
            docx_path,
        ]
        result = subprocess.run(
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            check=False,
            text=True,
        )
        # Generated PDF has the same stem as the source file
        generated_pdf = os.path.join(out_dir, Path(docx_path).with_suffix(".pdf").name)
        if not os.path.exists(generated_pdf):
            raise RuntimeError(
                f"LibreOffice conversion failed – see stdout/stderr.\n"
                f"stdout: {result.stdout}\nstderr: {result.stderr}"
            )
        # Move the file to the exact location requested by the caller
        shutil.move(generated_pdf, pdf_path)
        return True

    except Exception as e:  # pragma: no‑cover – any unexpected error
        st.error(f"❌ Failed to convert **{Path(docx_path).name}** – {e}")
        return False
    finally:
        # On Windows we need to un‑initialise COM for the current thread.
        if sys.platform.startswith("win"):
            try:
                import pythoncom
                pythoncom.CoUninitialize()
            except Exception:  # pragma: no‑cover
                pass


# --------------------------------------------------------------
# Page configuration & branding (KL Group)
# --------------------------------------------------------------
st.set_page_config(
    page_title="KL Tools",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.markdown(
    """
    <style>
        .main-header {font-size:2.5rem; font-weight:700; color:#FF6B00; text-align:center; margin-bottom:.5rem;}
        .sub-header {font-size:1.1rem; color:#1E3A8B; text-align:center; margin-bottom:2rem; font-weight:500;}
        .tool-card {background:#f8f9fa; border-radius:10px; padding:20px; border-left:5px solid #FF6B00; margin-bottom:20px;}
        .download-box {background:#fff7ed; border:2px solid #FF6B00; border-radius:10px; padding:20px; margin:20px 0;}
        .merge-field {background:#ffedd5; color:#c2410c; padding:2px 8px; border-radius:4px;
                       font-family:monospace; font-weight:600;}
        .debug-text {background:#f3f4f6; padding:10px; border-radius:4px; font-family:monospace;
                     font-size:0.9em; color:#374151; max-height:200px; overflow-y:auto;}
        .stButton>button {width:100%; border-radius:8px; height:3rem; font-weight:600;
                          background:#FF6B00; color:#fff; border:none;}
        .stButton>button:hover {background:#e55a00;}
        .footer-text {text-align:center; color:#1E3A8B; padding:20px; font-weight:600;}
        .kl-orange {color:#FF6B00;}
    </style>
    """,
    unsafe_allow_html=True,
)

# --------------------------------------------------------------
# Initialise session‑state keys (download persistence)
# --------------------------------------------------------------
for key in [
    "split_results", "split_zip", "split_single",
    "merge_result", "merge_filename", "merge_pages",
    "rename_results", "rename_zip", "rename_single",
    "map_results", "map_zip", "map_single",
    "mail_docs", "mail_zip", "mail_single", "mail_name", "mail_count",
    "word_zip", "word_single", "word_count",
    "excel_zip", "excel_single", "excel_count",
    "template_bytes",
    "png_files", "png_zip",
    "compress_zip", "compress_single", "compression_stats",
    # DOCX → PDF keys (renamed to avoid clashes with other tools)
    "docxpdf_zip", "docxpdf_single", "docxpdf_count",
]:
    if key not in st.session_state:
        st.session_state[key] = None

# --------------------------------------------------------------
# Header
# --------------------------------------------------------------
st.markdown('<div class="main-header">📄 KL Tools</div>', unsafe_allow_html=True)
st.markdown(
    '<div class="sub-header">Enhanced Document Management Suite by <span class="kl-orange">KL Group</span></div>',
    unsafe_allow_html=True,
)

# --------------------------------------------------------------
# Sidebar – tool selector & performance settings
# --------------------------------------------------------------
st.sidebar.markdown(
    '<div style="font-size:1.2rem; font-weight:600; color:#FF6B00; margin-bottom:1rem;">🛠️ Select Tool</div>',
    unsafe_allow_html=True,
)
tool = st.sidebar.radio(
    "",
    [
        "✂️ PDF Splitter",
        "🔗 PDF Merger",
        "🏷️ Payslip Renamer",
        "📧 Mail Merge",
        "📄 PDF to Word",
        "📄 DOCX to PDF",
        "📊 PDF to Excel",
        "🖼️ PDF to PNG",
        "📄 PNG to PDF",
        "📦 PDF Compressor",
    ],
    label_visibility="collapsed",
)
st.sidebar.markdown("---")
st.sidebar.markdown(
    """
    **💡 KL Tools Advantages**  
    • Secure local processing  
    • Bulk‑operation support  
    • Powered by **KL Group**
    """
)

# Slider for parallel workers (used by heavy‑batch tools)
max_workers = st.sidebar.slider(
    "Parallel workers (for batch operations)",
    min_value=1,
    max_value=os.cpu_count() or 4,
    value=4,
    help="Higher numbers use more CPU cores – useful for large batches.",
)

# --------------------------------------------------------------
# Helper utilities
# --------------------------------------------------------------
def create_zip_from_files(file_list):
    """file_list = [(filename, bytes_or_path), …] → ZIP binary."""
    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for name, data in file_list:
            if isinstance(data, (str, Path)):
                zf.write(data, name)
            else:
                zf.writestr(name, data)
    zip_buf.seek(0)
    return zip_buf.getvalue()


def extract_text_from_docx(doc):
    """Return all visible text in a DOCX (debug preview)."""
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for tbl in doc.tables:
        for row in tbl.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def detect_merge_fields(doc):
    """
    Detect merge fields in a DOCX.
    Supports {{Field}} (double‑curly) and «Field» (double‑angle‑quote).
    """
    fields = set()
    curly_pat = re.compile(r"\{\{([A-Za-z0-9_]+)\}\}")
    angle_pat = re.compile(r"«([^»]+)»")

    # Paragraphs (including runs)
    for para in doc.paragraphs:
        fields.update(curly_pat.findall(para.text))
        fields.update(angle_pat.findall(para.text))
        for run in para.runs:
            fields.update(curly_pat.findall(run.text))
            fields.update(angle_pat.findall(run.text))
        combined = "".join(run.text for run in para.runs)
        fields.update(curly_pat.findall(combined))
        fields.update(angle_pat.findall(combined))

    # Tables
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                fields.update(curly_pat.findall(cell.text))
                fields.update(angle_pat.findall(cell.text))
                for para in cell.paragraphs:
                    fields.update(curly_pat.findall(para.text))
                    fields.update(angle_pat.findall(para.text))
                    combined = "".join(run.text for run in para.runs)
                    fields.update(curly_pat.findall(combined))
                    fields.update(angle_pat.findall(combined))
    return sorted(fields)


def compress_pdf(input_bytes):
    """Lossless compression using pypdf (no external libs)."""
    try:
        reader = PdfReader(io.BytesIO(input_bytes))
        writer = PdfWriter()
        for page in reader.pages:
            page.compress_content_streams()
            writer.add_page(page)
        out_buf = io.BytesIO()
        writer.write(out_buf)
        return out_buf.getvalue()
    except Exception as e:
        st.error(f"❌ Compression error: {e}")
        return input_bytes


def parse_page_range(txt: str, max_page: int) -> list[int]:
    """Convert a page‑range string like “1,3‑5,8” into a sorted list of zero‑based page numbers."""
    pages = set()
    for part in txt.split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            start, end = part.split("-", 1)
            try:
                s = max(int(start) - 1, 0)
                e = min(int(end) - 1, max_page - 1)
                pages.update(range(s, e + 1))
            except ValueError:
                continue
        else:
            try:
                p = int(part) - 1
                if 0 <= p < max_page:
                    pages.add(p)
            except ValueError:
                continue
    return sorted(pages)


# --------------------------------------------------------------
# TOOL 1 – PDF Splitter (parallel per‑PDF)
# --------------------------------------------------------------
if tool == "✂️ PDF Splitter":
    st.markdown('<div class="tool-card">', unsafe_allow_html=True)
    st.subheader("✂️ PDF Splitter")
    st.markdown("Split PDFs into individual pages or custom sections.")
    st.markdown('</div>', unsafe_allow_html=True)

    left, right = st.columns(2)

    # ----- Upload -------------------------------------------------
    with left:
        pdf_files = st.file_uploader(
            "Drop PDF files (multiple allowed)", type=["pdf"], accept_multiple_files=True
        )
        if pdf_files:
            total_pages = sum(
                len(PdfReader(io.BytesIO(p.getvalue())).pages) for p in pdf_files
            )
            st.success(f"✅ {len(pdf_files)} file(s) – {total_pages} pages total")

    # ----- Settings ------------------------------------------------
    with right:
        if pdf_files:
            split_mode = st.selectbox(
                "Split mode",
                [
                    "1️⃣ One page → one file",
                    "🔢 Every N pages",
                    "🔎 By keyword",
                    "🛠️ Custom ranges",
                ],
            )
            cfg = {}
            if split_mode == "🔢 Every N pages":
                cfg["n"] = st.number_input("Pages per file", min_value=1, value=2)
            elif split_mode == "🔎 By keyword":
                cfg["keyword"] = st.text_input("Keyword that starts a new file")
            elif split_mode == "🛠️ Custom ranges":
                cfg["ranges"] = st.text_input(
                    "Ranges (e.g. 1,2,3 or 1-3,4-6)",
                    help="Comma‑separated single pages or start‑end pairs",
                )
            download_choice = st.radio(
                "Download as", ["📦 ZIP (individual)", "📄 Single combined PDF"]
            )

    # ----- Helper for a single PDF ---------------------------------
    def _process_one_pdf(pdf_file, mode, cfg_local):
        pdf_bytes = pdf_file.getvalue()
        reader = PdfReader(io.BytesIO(pdf_bytes))
        n_pages = len(reader.pages)
        base_name = Path(pdf_file.name).stem

        # ----- Determine ranges -----
        if mode == "1️⃣ One page → one file":
            ranges = [(i, i) for i in range(n_pages)]
        elif mode == "🔢 Every N pages":
            n = cfg_local["n"]
            ranges = [(i, min(i + n - 1, n_pages - 1)) for i in range(0, n_pages, n)]
        elif mode == "🔎 By keyword":
            kw = cfg_local["keyword"].lower()
            kw_pages = []
            with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdoc:
                for i, pg in enumerate(pdoc.pages):
                    txt = (pg.extract_text() or "").lower()
                    if kw in txt:
                        kw_pages.append(i)
            ranges = []
            for i, start in enumerate(kw_pages):
                end = kw_pages[i + 1] - 1 if i + 1 < len(kw_pages) else n_pages - 1
                ranges.append((start, end))
        else:  # Custom ranges
            ranges = []
            for part in cfg_local["ranges"].split(","):
                part = part.strip()
                if "-" in part:
                    s, e = part.split("-")
                    s = int(s) - 1
                    e = int(e) - 1 if e.strip().lower() not in ["end", "last"] else n_pages - 1
                    ranges.append((s, e))
                else:
                    p = int(part) - 1
                    ranges.append((p, p))

        # ----- Write each part to a temporary file -----
        out = []
        for i, (s, e) in enumerate(ranges):
            writer = PdfWriter()
            for p in range(s, e + 1):
                writer.add_page(reader.pages[p])
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
            writer.write(tmp)
            tmp.close()
            out_name = f"{base_name}_Part{i+1:03d}.pdf"
            out.append((out_name, tmp.name))
        return out

    # ----- Process -------------------------------------------------
    if pdf_files and st.button("🚀 Split PDFs", type="primary"):
        with st.spinner("Splitting PDFs (parallel)…"):
            progress = st.progress(0)
            all_parts = []   # [(filename, temp_path), …]

            with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
                futures = [
                    executor.submit(_process_one_pdf, pdf, split_mode, cfg)
                    for pdf in pdf_files
                ]
                for i, fut in enumerate(concurrent.futures.as_completed(futures), 1):
                    all_parts.extend(fut.result())
                    progress.progress(i / len(pdf_files))

            # ----- Build download artefact -----
            if download_choice == "📦 ZIP (individual)":
                st.session_state.split_zip = create_zip_from_files(all_parts)
                st.session_state.split_single = None
            else:
                merger = PdfWriter()
                for _, p in all_parts:
                    merger.append_pages_from_reader(PdfReader(p))
                merged_buf = io.BytesIO()
                merger.write(merged_buf)
                st.session_state.split_single = merged_buf.getvalue()
                st.session_state.split_zip = None

            # Cleanup temp files
            for _, p in all_parts:
                if os.path.exists(p):
                    os.remove(p)

            # Store simple result list for UI
            st.session_state.split_results = [{"File": fn} for fn, _ in all_parts]
            st.success(f"✅ Created {len(all_parts)} file(s)")

    # ----- Show results & download ---------------------------------
    if st.session_state.get("split_results"):
        st.markdown('<div class="download-box">', unsafe_allow_html=True)
        st.subheader("📥 Download")
        c1, c2 = st.columns(2)
        with c1:
            st.metric("Files created", len(st.session_state.split_results))
        with c2:
            if st.session_state.get("split_zip"):
                st.download_button("⬇️ ZIP", st.session_state.split_zip, "split_files.zip")
            else:
                st.download_button("⬇️ PDF", st.session_state.split_single, "combined.pdf")
        with st.expander("Details"):
            st.dataframe(pd.DataFrame(st.session_state.split_results))
        st.markdown("</div>", unsafe_allow_html=True)


# --------------------------------------------------------------
# TOOL 2 – PDF Merger (fast sequential)
# --------------------------------------------------------------
elif tool == "🔗 PDF Merger":
    st.markdown('<div class="tool-card">', unsafe_allow_html=True)
    st.subheader("🔗 PDF Merger")
    st.markdown("Combine multiple PDFs into a single document (optional bookmarks).")
    st.markdown('</div>', unsafe_allow_html=True)

    merge_files = st.file_uploader(
        "Upload PDFs in desired order (first = first pages)",
        type=["pdf"],
        accept_multiple_files=True,
        key="merge_upload",
    )
    if merge_files:
        st.success(f"✅ {len(merge_files)} files selected")
        add_bookmarks = st.checkbox("Add bookmarks (one per source file)", value=True)
        out_name = st.text_input("Output filename", "merged.pdf")
        if st.button("🔗 Merge PDFs", type="primary"):
            writer = PdfWriter()
            total_pages = 0
            page_counter = 0
            for pdf in merge_files:
                reader = PdfReader(io.BytesIO(pdf.getvalue()))
                writer.append_pages_from_reader(reader)
                if add_bookmarks:
                    writer.add_outline_item(Path(pdf.name).stem, page_counter)
                page_counter += len(reader.pages)
                total_pages += len(reader.pages)

            buf = io.BytesIO()
            writer.write(buf)
            st.session_state.merge_result = buf.getvalue()
            st.session_state.merge_filename = (
                out_name if out_name.lower().endswith(".pdf") else out_name + ".pdf"
            )
            st.session_state.merge_pages = total_pages
            st.success("✅ PDFs merged successfully")

    if st.session_state.get("merge_result"):
        st.markdown('<div class="download-box">', unsafe_allow_html=True)
        st.metric("Total pages", st.session_state.merge_pages)
        st.download_button(
            "⬇️ Download merged PDF",
            st.session_state.merge_result,
            st.session_state.merge_filename,
            "application/pdf",
        )
        st.markdown("</div>", unsafe_allow_html=True)


# --------------------------------------------------------------
# TOOL 3 – Payslip Renamer
# --------------------------------------------------------------
elif tool == "🏷️ Payslip Renamer":
    st.markdown('<div class="tool-card">', unsafe_allow_html=True)
    st.subheader("🏷️ Payslip Renamer")
    st.markdown(
        "Rename payslip PDFs by matching employee names (auto‑match) or by using an explicit Excel mapping."
    )
    st.markdown('</div>', unsafe_allow_html=True)

    tab_auto, tab_map = st.tabs(["🔍 Auto‑Match", "📋 Excel Mapping"])

    # ---------------- Auto‑Match ----------------
    with tab_auto:
        col_a, col_b = st.columns(2)
        with col_a:
            excel = st.file_uploader("Employee list (Excel)", type=["xlsx", "xls"], key="auto_excel")
            if excel:
                df_emp = pd.read_excel(excel)
                name_col = st.selectbox("Column with employee names", df_emp.columns)
                st.success(f"✅ {len(df_emp)} employees loaded")
        with col_b:
            pdfs = st.file_uploader("Payslip PDFs", type=["pdf"], accept_multiple_files=True, key="auto_pdfs")
            if pdfs:
                st.success(f"✅ {len(pdfs)} PDFs uploaded")

        if excel and pdfs:
            dl_choice = st.radio(
                "Download as", ["📦 ZIP (individual files)", "📄 Single combined PDF"], key="auto_dl"
            )
            if st.button("🚀 Rename", type="primary"):
                with st.spinner("Renaming…"):
                    employee_names = (
                        df_emp[name_col].dropna().astype(str).str.strip().tolist()
                    )
                    results = []
                    temp_files = []   # (new_name, temp_path)

                    for pdf in pdfs:
                        with pdfplumber.open(io.BytesIO(pdf.getvalue())) as doc:
                            txt = " ".join(p.extract_text() or "" for p in doc.pages)
                        best_match, best_score = None, 0
                        for emp in employee_names:
                            s = fuzz.partial_ratio(emp.lower(), txt.lower())
                            if s > best_score and s > 80:
                                best_match, best_score = emp, s
                        new_name = (
                            re.sub(r'[<>:"/\\|?*]', "", best_match) + "_Payslip.pdf"
                            if best_match
                            else f"Unmatched_{pdf.name}"
                        )
                        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
                        tmp.write(pdf.getvalue())
                        tmp.close()
                        temp_files.append((new_name, tmp.name))
                        results.append(
                            {"Original": pdf.name, "New": new_name, "Match": best_match or "—"}
                        )

                    if dl_choice == "📦 ZIP (individual files)":
                        st.session_state.rename_zip = create_zip_from_files(temp_files)
                        st.session_state.rename_single = None
                    else:
                        merger = PdfWriter()
                        for _, p in temp_files:
                            merger.append_pages_from_reader(PdfReader(p))
                        merged_buf = io.BytesIO()
                        merger.write(merged_buf)
                        st.session_state.rename_single = merged_buf.getvalue()
                        st.session_state.rename_zip = None

                    st.session_state.rename_results = results
                    st.success(f"✅ Processed {len(results)} PDFs")

        if st.session_state.get("rename_results"):
            st.markdown('<div class="download-box">', unsafe_allow_html=True)
            st.subheader("📥 Download")
            c1, c2 = st.columns(2)
            with c1:
                matched = sum(
                    1 for r in st.session_state.rename_results if r["Match"] != "—"
                )
                st.metric("Matched", f"{matched}/{len(st.session_state.rename_results)}")
            with c2:
                if st.session_state.get("rename_zip"):
                    st.download_button("⬇️ ZIP", st.session_state.rename_zip, "renamed.zip")
                else:
                    st.download_button("⬇️ PDF", st.session_state.rename_single, "all_payslips.pdf")
            with st.expander("Details"):
                st.dataframe(pd.DataFrame(st.session_state.rename_results))
            st.markdown("</div>", unsafe_allow_html=True)

    # ---------------- Excel Mapping ----------------
    with tab_map:
        map_excel = st.file_uploader(
            "Mapping Excel (Current → New)", type=["xlsx", "xls"], key="map_excel"
        )
        map_pdfs = st.file_uploader(
            "Payslip PDFs to rename", type=["pdf"], accept_multiple_files=True, key="map_pdfs"
        )
        if map_excel and map_pdfs:
            df_map = pd.read_excel(map_excel)
            cur_col = st.selectbox("Current filename column", df_map.columns, key="cur_col")
            new_col = st.selectbox("New name column", df_map.columns, key="new_col")
            dl_choice = st.radio("Download as", ["📦 ZIP", "📄 Single combined PDF"], key="map_dl")
            if st.button("🚀 Apply mapping", type="primary"):
                name_dict = dict(zip(df_map[cur_col].astype(str), df_map[new_col].astype(str)))
                results = []
                temp_files = []   # (new_name, temp_path)

                for pdf in map_pdfs:
                    new_name = (
                        re.sub(r'[<>:"/\\|?*]', "", name_dict.get(pdf.name, f"Unmatched_{pdf.name}"))
                        + ".pdf"
                    )
                    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
                    tmp.write(pdf.getvalue())
                    tmp.close()
                    temp_files.append((new_name, tmp.name))
                    results.append({"Original": pdf.name, "New": new_name})

                if dl_choice == "📦 ZIP":
                    st.session_state.map_zip = create_zip_from_files(temp_files)
                    st.session_state.map_single = None
                else:
                    merger = PdfWriter()
                    for _, p in temp_files:
                        merger.append_pages_from_reader(PdfReader(p))
                    merged_buf = io.BytesIO()
                    merger.write(merged_buf)
                    st.session_state.map_single = merged_buf.getvalue()
                    st.session_state.map_zip = None

                st.session_state.map_results = results
                st.success("✅ Mapping applied")

        if st.session_state.get("map_results"):
            st.markdown('<div class="download-box">', unsafe_allow_html=True)
            c1, c2 = st.columns(2)
            with c1:
                st.metric("Files processed", len(st.session_state.map_results))
            with c2:
                if st.session_state.get("map_zip"):
                    st.download_button("⬇️ ZIP", st.session_state.map_zip, "mapped.zip")
                else:
                    st.download_button("⬇️ PDF", st.session_state.map_single, "mapped.pdf")
            with st.expander("Details"):
                st.dataframe(pd.DataFrame(st.session_state.map_results))
            st.markdown("</div>", unsafe_allow_html=True)


# --------------------------------------------------------------
# TOOL 4 – Enhanced Mail Merge (Word & PDF output, parallelised)
# --------------------------------------------------------------
elif tool == "📧 Mail Merge":
    st.markdown('<div class="tool-card">', unsafe_allow_html=True)
    st.subheader("📧 Mail Merge")
    st.markdown(
        "Create bulk personalized documents from a template and an Excel sheet.<br>"
        "<b>Supported placeholders:</b> <span class='merge-field'>{{Field}}</span> or <span class='merge-field'>«Field»</span>",
        unsafe_allow_html=True,
    )
    st.markdown('</div>', unsafe_allow_html=True)

    # ----- 1️⃣ Upload template ------------------------------------
    tmpl_file = st.file_uploader("Word template (.docx)", type=["docx"], key="mail_template")
    template_bytes = None
    merge_fields = []

    if tmpl_file:
        template_bytes = tmpl_file.getvalue()
        st.session_state.template_bytes = template_bytes
        doc = Document(io.BytesIO(template_bytes))
        merge_fields = detect_merge_fields(doc)

        with st.expander("🔍 Template preview"):
            txt = extract_text_from_docx(doc)
            preview = txt[:800] + ("…" if len(txt) > 800 else "")
            st.markdown(
                f"<div class='debug-text'>{preview.replace('<', '&lt;').replace('>', '&gt;')}</div>",
                unsafe_allow_html=True,
            )

        if merge_fields:
            placeholder_str = ", ".join([f"{{{{{fld}}}}}" for fld in merge_fields])
            st.success(f"✅ Detected placeholders: {placeholder_str}")
        else:
            st.error("❌ No placeholders detected!")
            manual = st.text_input("Enter placeholders manually (comma‑separated)", "Emp_Id,Emp_Name,Basic")
            if manual:
                merge_fields = [f.strip() for f in manual.split(",") if f.strip()]

    # ----- 2️⃣ Upload Excel data -----------------------------------
    data_file = st.file_uploader("Excel data", type=["xlsx", "xls"], key="mail_excel")
    df = None
    if data_file:
        df = pd.read_excel(data_file)
        st.success(f"✅ {len(df)} rows – columns: {', '.join(df.columns)}")
        with st.expander("Data preview"):
            st.dataframe(df.head())

    # ----- 3️⃣ Mapping & generation (parallel) --------------------
    if template_bytes and df is not None and merge_fields:
        st.markdown("### 🗂️ Map placeholders to Excel columns")
        field_map = {}
        cols = st.columns(min(len(merge_fields), 3))
        for i, fld in enumerate(merge_fields):
            with cols[i % len(cols)]:
                default_idx = 0
                for idx, col in enumerate(df.columns):
                    if (
                        fld.lower() == col.lower()
                        or fld.lower().replace("_", "") == col.lower().replace(" ", "")
                    ):
                        default_idx = idx
                        break
                field_map[fld] = st.selectbox(
                    f"{{{{{fld}}}}}", df.columns, index=default_idx, key=f"map_{fld}"
                )

        st.markdown("### ⚙️ Output settings")
        col_out, col_name = st.columns(2)
        with col_out:
            output_format = st.radio(
                "Output format", ["📄 Word (.docx)", "📊 PDF (.pdf)"], key="output_format"
            )
        with col_name:
            naming_field = st.selectbox(
                "Name files by column", list(field_map.values()), key="naming_field"
            )
        dl_format = st.radio(
            "Download as", ["📦 ZIP (individual)", "🔄 Single file"], key="dl_format"
        )

        if st.button("🚀 Generate documents", type="primary"):
            with st.spinner(f"Generating {len(df)} documents (parallel)…"):
                generated = []   # [(filename, bytes), …]

                def _create_one(idx, row):
                    doc = Document(io.BytesIO(template_bytes))

                    # Replace placeholders in paragraphs
                    for para in doc.paragraphs:
                        combined = "".join(run.text for run in para.runs)
                        new_text = combined
                        for fld, col in field_map.items():
                            val = str(row.get(col, ""))
                            new_text = new_text.replace(f"{{{{{fld}}}}}", val).replace(f"«{fld}»", val)
                        if new_text != combined:
                            para.clear()
                            para.add_run(new_text)

                    # Replace inside tables
                    for table in doc.tables:
                        for row_obj in table.rows:
                            for cell in row_obj.cells:
                                for para in cell.paragraphs:
                                    combined = "".join(run.text for run in para.runs)
                                    new_text = combined
                                    for fld, col in field_map.items():
                                        val = str(row.get(col, ""))
                                        new_text = new_text.replace(f"{{{{{fld}}}}}", val).replace(f"«{fld}»", val)
                                    if new_text != combined:
                                        para.clear()
                                        para.add_run(new_text)

                    # Save according to chosen format
                    if output_format == "📄 Word (.docx)":
                        out_buf = io.BytesIO()
                        doc.save(out_buf)
                        file_bytes = out_buf.getvalue()
                        ext = ".docx"
                    else:
                        # DOCX → PDF (fast, with COM init)
                        tmp_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
                        doc.save(tmp_docx.name)
                        tmp_docx.close()
                        tmp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
                        tmp_pdf.close()
                        success = convert_docx_to_pdf(tmp_docx.name, tmp_pdf.name)
                        if not success:
                            os.remove(tmp_docx.name)
                            os.remove(tmp_pdf.name)
                            return None
                        with open(tmp_pdf.name, "rb") as f:
                            file_bytes = f.read()
                        os.remove(tmp_docx.name)
                        os.remove(tmp_pdf.name)
                        ext = ".pdf"

                    safe_name = re.sub(r'[<>:"/\\|?*]', "", str(row.get(naming_field, f"Doc_{idx+1}")))[:50]
                    filename = f"{safe_name}{ext}"
                    return filename, file_bytes

                with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
                    futures = [
                        executor.submit(_create_one, i, row)
                        for i, row in enumerate(df.to_dict("records"))
                    ]
                    for fut in concurrent.futures.as_completed(futures):
                        res = fut.result()
                        if res:
                            generated.append(res)

                # ----- Build final artefact -----
                if dl_format == "📦 ZIP (individual)":
                    st.session_state.mail_zip = create_zip_from_files(generated)
                    st.session_state.mail_single = None
                else:
                    if output_format == "📄 Word (.docx)":
                        st.session_state.mail_single = generated[0][1] if generated else None
                    else:
                        merger = PdfWriter()
                        for _, data in generated:
                            merger.append_pages_from_reader(PdfReader(io.BytesIO(data)))
                        merged_buf = io.BytesIO()
                        merger.write(merged_buf)
                        st.session_state.mail_single = merged_buf.getvalue()
                    st.session_state.mail_zip = None

                st.session_state.mail_count = len(generated)
                st.success(f"✅ Generated {len(generated)} document(s)")

    # ----- Download area -------------------------------------------
    if st.session_state.get("mail_count"):
        st.markdown('<div class="download-box">', unsafe_allow_html=True)
        st.subheader("📥 Download results")
        c1, c2 = st.columns(2)
        with c1:
            st.metric("Files", st.session_state.mail_count)
        with c2:
            if st.session_state.get("mail_zip"):
                st.download_button("⬇️ ZIP", st.session_state.mail_zip, "mail_merge.zip")
            elif st.session_state.get("mail_single"):
                ext = ".pdf" if output_format == "📊 PDF (.pdf)" else ".docx"
                mime = "application/pdf" if output_format == "📊 PDF (.pdf)" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                st.download_button(f"⬇️ Download{ext}", st.session_state.mail_single, f"merged{ext}", mime)
        st.markdown("</div>", unsafe_allow_html=True)


# --------------------------------------------------------------
# TOOL 5 – PDF → Word (parallelised)
# --------------------------------------------------------------
elif tool == "📄 PDF to Word":
    st.markdown('<div class="tool-card">', unsafe_allow_html=True)
    st.subheader("📄 PDF → Word")
    st.markdown("Convert PDFs into editable DOCX files.")
    st.markdown('</div>', unsafe_allow_html=True)

    pdfs = st.file_uploader(
        "Select PDFs", type=["pdf"], accept_multiple_files=True, key="pdf2word"
    )
    if pdfs:
        out_mode = st.radio(
            "Download as", ["📦 ZIP (individual DOCX)", "📄 Single DOCX (first file only)"]
        )
        if st.button("📄 Convert", type="primary"):
            files = []

            def _pdf_to_docx_one(pdf):
                pdf_path = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
                pdf_path.write(pdf.getvalue())
                pdf_path.close()
                docx_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
                docx_path.close()
                try:
                    conv = Converter(pdf_path.name)
                    conv.convert(docx_path.name)
                    conv.close()
                    with open(docx_path.name, "rb") as f:
                        data = f.read()
                    return (Path(pdf.name).stem + ".docx", data)
                finally:
                    os.remove(pdf_path.name)
                    os.remove(docx_path.name)

            with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
                futures = [executor.submit(_pdf_to_docx_one, pdf) for pdf in pdfs]
                for fut in concurrent.futures.as_completed(futures):
                    files.append(fut.result())

            if out_mode == "📦 ZIP (individual DOCX)":
                st.session_state.word_zip = create_zip_from_files(files)
                st.session_state.word_single = None
            else:
                st.session_state.word_single = files[0][1] if files else None
                st.session_state.word_zip = None
            st.session_state.word_count = len(files)
            st.success(f"✅ Converted {len(files)} file(s)")

    if st.session_state.get("word_count"):
        st.markdown('<div class="download-box">', unsafe_allow_html=True)
        if st.session_state.get("word_zip"):
            st.download_button("⬇️ ZIP", st.session_state.word_zip, "word_files.zip")
        else:
            st.download_button(
                "⬇️ DOCX",
                st.session_state.word_single,
                "converted.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        st.markdown("</div>", unsafe_allow_html=True)


# --------------------------------------------------------------
# TOOL 6 – DOCX → PDF (fast, parallel)
# --------------------------------------------------------------
elif tool == "📄 DOCX to PDF":
    st.markdown('<div class="tool-card">', unsafe_allow_html=True)
    st.subheader("📄 DOCX → PDF")
    st.markdown("Convert Word documents to PDF while preserving layout.")
    st.markdown('</div>', unsafe_allow_html=True)

    # ---- File upload -------------------------------------------------
    docx_files = st.file_uploader(
        "Select DOCX files", type=["docx"], accept_multiple_files=True, key="docx2pdf"
    )
    if docx_files:
        out_mode = st.radio(
            "Download as", ["📦 ZIP (individual PDFs)", "📄 Single merged PDF"],
            key="docx2pdf_out"
        )
        if st.checkbox("Show preview of first document (debug)"):
            doc = Document(io.BytesIO(docx_files[0].getvalue()))
            preview_txt = extract_text_from_docx(doc)
            st.code(preview_txt[:2000] + ("…" if len(preview_txt) > 2000 else ""))

        if st.button("📄 Convert to PDF", type="primary"):
            with st.spinner("Converting DOCX files to PDF (parallel)…"):
                pdfs = []   # list of (filename, bytes)

                def _docx_to_pdf_one(docx):
                    # Write uploaded docx to a temp file
                    tmp_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
                    tmp_docx.write(docx.getvalue())
                    tmp_docx.close()

                    # Destination PDF (another temp file)
                    tmp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
                    tmp_pdf.close()

                    # ----- Actual conversion -----
                    success = convert_docx_to_pdf(tmp_docx.name, tmp_pdf.name)
                    if not success:
                        # Clean up any leftovers and return None → filtered out later
                        os.remove(tmp_docx.name)
                        os.remove(tmp_pdf.name)
                        return None

                    # Read the generated PDF back into memory
                    with open(tmp_pdf.name, "rb") as f:
                        data = f.read()

                    # Remove temporary files
                    os.remove(tmp_docx.name)
                    os.remove(tmp_pdf.name)

                    # Return a user‑friendly name (original stem + .pdf)
                    return (Path(docx.name).stem + ".pdf", data)

                # ---- Parallel execution ----
                with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
                    futures = [executor.submit(_docx_to_pdf_one, d) for d in docx_files]
                    for fut in concurrent.futures.as_completed(futures):
                        res = fut.result()
                        if res:
                            pdfs.append(res)

                # ---- Build the download artefact ----
                if out_mode == "📦 ZIP (individual PDFs)":
                    st.session_state.docxpdf_zip = create_zip_from_files(pdfs)
                    st.session_state.docxpdf_single = None
                else:
                    # Merge everything into one PDF
                    merger = PdfWriter()
                    for _, data in pdfs:
                        merger.append_pages_from_reader(PdfReader(io.BytesIO(data)))
                    merged_buf = io.BytesIO()
                    merger.write(merged_buf)
                    st.session_state.docxpdf_single = merged_buf.getvalue()
                    st.session_state.docxpdf_zip = None

                st.session_state.docxpdf_count = len(pdfs)
                st.success(f"✅ Converted {len(pdfs)} document(s)")

    if st.session_state.get("docxpdf_count"):
        st.markdown('<div class="download-box">', unsafe_allow_html=True)
        st.subheader("📥 Download PDFs")
        c1, c2 = st.columns(2)
        with c1:
            st.metric("PDFs created", st.session_state.docxpdf_count)
        with c2:
            if st.session_state.get("docxpdf_zip"):
                st.download_button(
                    "⬇️ ZIP",
                    st.session_state.docxpdf_zip,
                    "docx_to_pdf.zip",
                    mime="application/zip",
                )
            else:
                st.download_button(
                    "⬇️ PDF",
                    st.session_state.docxpdf_single,
                    "merged.pdf",
                    mime="application/pdf",
                )
        st.markdown("</div>", unsafe_allow_html=True)


# --------------------------------------------------------------
# TOOL 7 – PDF → Excel (fast, typo fixed)
# --------------------------------------------------------------
elif tool == "📊 PDF to Excel":
    st.markdown('<div class="tool-card">', unsafe_allow_html=True)
    st.subheader("📊 PDF → Excel")
    st.markdown("Extract tables from PDFs into Excel workbooks.")
    st.markdown('</div>', unsafe_allow_html=True)

    pdfs = st.file_uploader(
        "Select PDFs", type=["pdf"], accept_multiple_files=True, key="pdf2excel"
    )
    if pdfs:
        mode = st.radio("Extract from", ["All pages", "Specific pages"])
        pages_input = ""
        if mode == "Specific pages":
            pages_input = st.text_input("Pages (comma‑separated)", "1,2,3")
        dl_mode = st.radio(
            "Download as", ["📦 ZIP (one Excel per PDF)", "📄 Combined workbook"]
        )
        if st.button("📊 Extract tables", type="primary"):
            excel_files = []  # [(filename, bytes), …]

            for pdf in pdfs:
                with pdfplumber.open(io.BytesIO(pdf.getvalue())) as doc:
                    if mode == "Specific pages":
                        page_nums = parse_page_range(pages_input, len(doc.pages))
                    else:
                        page_nums = list(range(len(doc.pages)))

                    sheets = {}
                    for pnum in page_nums:
                        if pnum >= len(doc.pages):
                            continue
                        tables = doc.pages[pnum].extract_tables()
                        for t_idx, table in enumerate(tables):
                            if not table:
                                continue
                            header = table[0] if any(cell for cell in table[0]) else None
                            df_tab = pd.DataFrame(table[1:], columns=header)
                            sheets[f"P{pnum+1}_T{t_idx+1}"] = df_tab

                    if sheets:
                        buf = io.BytesIO()
                        with pd.ExcelWriter(buf, engine="openpyxl") as writer:
                            for sheet_name, df_tab in sheets.items():
                                df_tab.to_excel(writer, sheet_name=sheet_name[:31], index=False)
                        excel_files.append((Path(pdf.name).stem + ".xlsx", buf.getvalue()))
                    else:
                        st.info(f"ℹ️ No tables detected in **{pdf.name}**")

            if dl_mode == "📦 ZIP (one Excel per PDF)":
                st.session_state.excel_zip = create_zip_from_files(excel_files)
                st.session_state.excel_single = None
            else:
                combo_buf = io.BytesIO()
                with pd.ExcelWriter(combo_buf, engine="openpyxl") as writer:
                    for fname, data in excel_files:
                        xl = pd.ExcelFile(io.BytesIO(data))
                        for sh in xl.sheet_names:
                            df_sh = pd.read_excel(io.BytesIO(data), sheet_name=sh)
                            new_name = f"{Path(fname).stem}_{sh}"[:31]
                            df_sh.to_excel(writer, sheet_name=new_name, index=False)
                combo_buf.seek(0)
                st.session_state.excel_single = combo_buf.getvalue()
                st.session_state.excel_zip = None

            st.session_state.excel_count = len(excel_files)
            st.success(f"✅ Extracted tables from {len(excel_files)} PDF(s)")

    if st.session_state.get("excel_count"):
        st.markdown('<div class="download-box">', unsafe_allow_html=True)
        if st.session_state.get("excel_zip"):
            st.download_button(
                "⬇️ ZIP", st.session_state.excel_zip, "excel_files.zip", "application/zip"
            )
        else:
            st.download_button(
                "⬇️ Excel",
                st.session_state.excel_single,
                "combined_tables.xlsx",
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
        st.markdown("</div>", unsafe_allow_html=True)


# --------------------------------------------------------------
# TOOL 8 – PDF → PNG
# --------------------------------------------------------------
elif tool == "🖼️ PDF to PNG":
    st.markdown('<div class="tool-card">', unsafe_allow_html=True)
    st.subheader("🖼️ PDF → PNG")
    st.markdown("Convert PDF pages to high‑quality PNG images.")
    st.markdown('</div>', unsafe_allow_html=True)

    pdf_files = st.file_uploader(
        "Select PDF files", type=["pdf"], accept_multiple_files=True, key="pdf2png"
    )
    if pdf_files:
        col1, col2 = st.columns(2)
        with col1:
            dpi = st.slider(
                "Image quality (DPI)", min_value=72, max_value=300, value=150,
                help="Higher DPI = better quality but larger files"
            )
        with col2:
            download_mode = st.radio(
                "Download as",
                ["📦 ZIP (individual PNGs)", "📑 Single PDF (all images)"],
            )
        if st.button("🖼️ Convert to PNG", type="primary"):
            with st.spinner("Converting PDF pages to PNG…"):
                png_files = []  # [(filename, bytes), …]

                for pdf_file in pdf_files:
                    pdf_bytes = pdf_file.getvalue()
                    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                        for page_num in range(len(pdf.pages)):
                            page = pdf.pages[page_num]
                            pil_img = page.to_image(resolution=dpi).original
                            img_buf = io.BytesIO()
                            pil_img.save(img_buf, format="PNG", dpi=(dpi, dpi))
                            img_buf.seek(0)
                            fname = f"{Path(pdf_file.name).stem}_page_{page_num+1:03d}.png"
                            png_files.append((fname, img_buf.getvalue()))

                if download_mode == "📦 ZIP (individual PNGs)":
                    st.session_state.png_zip = create_zip_from_files(png_files)
                    st.success(f"✅ Converted {len(png_files)} page(s) to PNG")
                else:
                    # Combine PNGs into a single PDF (lossless via Pillow)
                    img_objs = [Image.open(io.BytesIO(b)) for _, b in png_files]
                    pdf_buf = io.BytesIO()
                    if img_objs:
                        img_objs[0].save(pdf_buf, format="PDF", save_all=True, append_images=img_objs[1:])
                    pdf_buf.seek(0)
                    # Return as ZIP for UI consistency
                    st.session_state.png_zip = create_zip_from_files([("images.pdf", pdf_buf.getvalue())])
                    st.success("✅ PNG pages combined into a single PDF (inside ZIP)")

                st.session_state.png_files = png_files

    if st.session_state.get("png_files"):
        st.markdown('<div class="download-box">', unsafe_allow_html=True)
        st.subheader("📥 Download PNG images")
        if st.session_state.get("png_zip"):
            st.download_button(
                "⬇️ Download ZIP",
                st.session_state.png_zip,
                "pdf_to_png.zip",
                "application/zip",
            )
        st.markdown("</div>", unsafe_allow_html=True)


# --------------------------------------------------------------
# TOOL 9 – PNG → PDF (pixel‑perfect)
# --------------------------------------------------------------
elif tool == "📄 PNG to PDF":
    st.markdown('<div class="tool-card">', unsafe_allow_html=True)
    st.subheader("📄 PNG → PDF")
    st.markdown("Convert PNG/JPG images into a single PDF document (pixel‑perfect).")
    st.markdown('</div>', unsafe_allow_html=True)

    img_files = st.file_uploader(
        "Select PNG/JPG images", type=["png", "jpg", "jpeg"], accept_multiple_files=True, key="png2pdf"
    )
    if img_files:
        if st.button("📄 Convert to PDF", type="primary"):
            with st.spinner("Creating PDF…"):
                image_bytes = [img.read() for img in img_files]
                try:
                    pdf_bytes = img2pdf.convert(image_bytes)      # loss‑less
                    st.session_state.png_to_pdf_result = pdf_bytes
                    st.success(f"✅ PDF created with {len(image_bytes)} page(s)")
                except Exception as e:
                    st.error(f"❌ Failed to create PDF: {e}")

    if st.session_state.get("png_to_pdf_result"):
        st.markdown('<div class="download-box">', unsafe_allow_html=True)
        st.download_button(
            "⬇️ Download PDF",
            st.session_state.png_to_pdf_result,
            "images_combined.pdf",
            "application/pdf",
        )
        st.markdown("</div>", unsafe_allow_html=True)


# --------------------------------------------------------------
# TOOL 10 – PDF Compressor (fixed “Page must be part of a PdfWriter”)
# --------------------------------------------------------------
elif tool == "📦 PDF Compressor":
    st.markdown('<div class="tool-card">', unsafe_allow_html=True)
    st.subheader("📦 PDF Compressor")
    st.markdown("Reduce PDF file size without losing quality.")
    st.markdown('</div>', unsafe_allow_html=True)

    pdf_files = st.file_uploader(
        "Select PDF files to compress", type=["pdf"], accept_multiple_files=True, key="pdf_compress"
    )
    if pdf_files:
        col1, col2 = st.columns(2)
        with col1:
            compression_level = st.selectbox(
                "Compression level",
                ["Light (fast)", "Medium", "High (slower)"],
                help="Higher compression = smaller files but longer processing time",
            )
        with col2:
            download_mode = st.radio(
                "Download as", ["📦 ZIP (individual)", "📄 Single compressed PDF"]
            )
        if st.button("📦 Compress PDFs", type="primary"):
            with st.spinner("Compressing PDFs (parallel)…"):
                def _compress_one(pdf):
                    orig = pdf.getvalue()
                    comp = compress_pdf(orig)
                    return (f"compressed_{pdf.name}", comp, len(orig), len(comp))

                compressed = []
                original_total = 0
                compressed_total = 0

                with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
                    futures = [executor.submit(_compress_one, p) for p in pdf_files]
                    for fut in concurrent.futures.as_completed(futures):
                        name, data, orig_len, comp_len = fut.result()
                        compressed.append((name, data))
                        original_total += orig_len
                        compressed_total += comp_len

                if download_mode == "📦 ZIP (individual)":
                    st.session_state.compress_zip = create_zip_from_files(compressed)
                    st.session_state.compress_single = None
                else:
                    merger = PdfWriter()
                    for _, data in compressed:
                        merger.append_pages_from_reader(PdfReader(io.BytesIO(data)))
                    merged_buf = io.BytesIO()
                    merger.write(merged_buf)
                    st.session_state.compress_single = merged_buf.getvalue()
                    st.session_state.compress_zip = None

                savings = (
                    (original_total - compressed_total) / original_total * 100
                    if original_total > 0
                    else 0
                )
                st.session_state.compression_stats = {
                    "original_size": original_total,
                    "compressed_size": compressed_total,
                    "savings": savings,
                    "file_count": len(compressed),
                }
                st.success(
                    f"✅ Compression saved {savings:.1f}% across {len(compressed)} file(s)"
                )

    if st.session_state.get("compression_stats"):
        stats = st.session_state.compression_stats
        st.markdown('<div class="download-box">', unsafe_allow_html=True)
        st.subheader("📥 Download compressed files")
        c1, c2, c3 = st.columns(3)
        with c1:
            st.metric("Original size", f"{stats['original_size']/1024:.0f} KB")
        with c2:
            st.metric("Compressed size", f"{stats['compressed_size']/1024:.0f} KB")
        with c3:
            st.metric("Saved", f"{stats['savings']:.1f}%")
        if st.session_state.get("compress_zip"):
            st.download_button(
                "⬇️ Download ZIP", st.session_state.compress_zip, "compressed_pdfs.zip"
            )
        elif st.session_state.get("compress_single"):
            st.download_button(
                "⬇️ Download PDF",
                st.session_state.compress_single,
                "compressed_merged.pdf",
                "application/pdf",
            )
        st.markdown("</div>", unsafe_allow_html=True)


# --------------------------------------------------------------
# Footer
# --------------------------------------------------------------
st.markdown("---")
st.markdown(
    '<div class="footer-text"><span class="kl-orange">KL Tools</span> • Enhanced Document Management • Powered by <span class="kl-orange">KL Group</span></div>',
    unsafe_allow_html=True,
)
